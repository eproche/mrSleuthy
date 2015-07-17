import click
import nltk, glob, os, sys, getopt
from gensim.models import word2vec
from nltk.stem.snowball import SnowballStemmer
from docx import Document
from collections import Counter 
import logging
from PIL import Image, ImageDraw, ImageColor
import numpy as np
from numpy import *
from scipy.spatial.distance import cosine
import networkx as nx
import matplotlib.pyplot as plt
from sklearn import manifold
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer
import cPickle as pickle
import csv
import scipy.cluster.hierarchy as hc 
import pandas as pd 
import py
import simplejson as json
import base64

logging.basicConfig(format='%(asctime)s : %(levelname)s : %(message)s',\
    level=logging.INFO)

#If there's a different set of pre-trained entity vectors that you would like to use then load them as "model"
#Comment out the line below if you've already run word2vec for a specific input, the results should be pickled.
#model = word2vec.Word2Vec.load_word2vec_format('/Users/epr397/Downloads/GoogleNews-vectors-negative300.bin', binary=True) 


class Config(dict):  
    def __init__(self, *args, **kwargs):
        self.config = py.path.local(
            click.get_app_dir('sleuthing')
        ).join('config.json') # A

        super(Config, self).__init__(*args, **kwargs)

    def load(self):
        """load a JSON config file from disk"""
        try:
            self.update(json.loads(self.config.read())) # B
        except py.error.ENOENT:
            pass

    def save(self):
        self.config.ensure()
        with self.config.open('w') as f: # B
            f.write(json.dumps(self))

class NumpyEncoder(json.JSONEncoder):
    def default(self, obj):
        """
        if input object is a ndarray it will be converted into a dict holding dtype, shape and the data base64 encoded
        """
        if isinstance(obj, np.ndarray):
            data_b64 = base64.b64encode(obj.data)
            return dict(__ndarray__=data_b64,
                        dtype=str(obj.dtype),
                        shape=obj.shape)
        # Let the base class default method raise the TypeError
        return json.JSONEncoder(self, obj)


def json_numpy(dct):
    """
    Decodes a previously encoded numpy ndarray with proper shape and dtype
    :param dct: (dict) json encoded ndarray
    :return: (ndarray) if input was an encoded ndarray
    """
    if isinstance(dct, dict) and '__ndarray__' in dct:
        data = base64.b64decode(dct['__ndarray__'])
        return np.frombuffer(data, dct['dtype']).reshape(dct['shape'])
    return dct

pass_config = click.make_pass_decorator(Config, ensure=True)

@click.group()
@pass_config
def cli(config):
	"""To start Mr.Sleuthy, run 'start readin INPUT', where input is the filepath you want to read inAfter 'start' has run successfully, run 'word2vec' or 'tfidf' to compute a similarity matrix
	'Output' is used after 'tfidf' or 'word2vec' to generate visual outputsTo view options for commands, enter 'start command --help' """
	config.load()
	config.save()

@cli.command('readin', short_help= click.style('Read an input file', fg='blue'))
@click.option('-t', is_flag=True, default = False, help = click.style("Use a folder of .txt files as the input", fg='cyan'))
@click.option('--des', is_flag=True, default = True, help = click.style("Only use descriptions responses", fg='blue'))
@click.option('--rem', is_flag=True, default = True, help = click.style("Only use reminded responses", fg='cyan'))
@click.option('--stem', is_flag=True, default=False, help = click.style('Only use with tfidf command! Stem words with NLTK SnowballStemmer', fg='blue'))
@click.argument('inp', default='', required = True)
@pass_config
def readin(config, inp, des, rem, t, stem):
	"""Optional flags can be used to alter how input is read
	--stem is intended to be combined with the 'tfidf' command."""
	if t:
		output = read_folder(inp) 
	else:
		output = read_csv(inp) 
	images = output[0]
	answers = output[1]
	config['answers'] = answers
	config['images'] = images		
	if stem:
		config['storage'] = read_words(answers, rem, des, True)
	else: 
		config['storage'] = read_words(answers, rem, des, False)
	config['inp'] = inp
	config.save()
	click.secho("\ninput has been read, ready for analysis", fg='green')
	click.secho("run tfidf or word2vec to generate similarity matrix", fg='blue', blink=False)

@cli.command('write', short_help = click.style('Write responses, usable words, and POS counts', fg='cyan'))
@pass_config
def write(config):
	"""Write responses, usable words, and POS counts after readin"""
	clean_write(config.get('answers'))

@cli.command('tfidf', short_help = click.style('Calculate similarites with tf-idf', fg='green'))
@pass_config
def tfidf(config):
	"""Welcome to the tfidf zone (fast)"""
	tf = TfidfVectorizer()
	tf_mat = tf.fit_transform(config.get('storage'))
	click.secho('\ntfidf matrix generated\n', fg='green')
	sim_mat = cosine_similarity(tf_mat, tf_mat)
	json_mat = json.dumps(sim_mat, cls=NumpyEncoder)
	config['sim_mat'] = json_mat
	config.save()
	click.secho('saved tfidf sim_mat', fg='green')
	click.secho('Now run output to generate visualizations!', fg='red', blink=False)
		
@cli.command('word2vec', short_help = click.style("Calculate similarites with word2vec model", fg='yellow'))
@click.option('--unweighted', is_flag=True, default =False, help = click.style('Use with word2vec if you do not want tf-idf weightings attached', fg='yellow'))
@pass_config
def word2vec(config, unweighted):
	"""Welcome to the word2vec zone (slow) Run word2vec after a file has been read in"""
	inp = config.get('inp')
	inp2 = inp[4:]  #take off the 'csv/' folder from input string
	if unweighted:
		pickle_filepath = "pickles/pickle_"+str(inp2)+"_unweighted"
	else:
		pickle_filepath = "pickles/pickle_"+str(inp2)
	if os.path.exists(pickle_filepath): #
		click.secho("\ngrabbing "+pickle_filepath+" from the pickle jar", fg='green')
		with open(pickle_filepath) as pickle_handle:
			sim_mat = pickle.load(pickle_handle) #load sim_mat if it's already been generated
	else:
		if unweighted == False: # use tf-idf weights
			tf = TfidfVectorizer()
			tf_mat = tf.fit_transform(config.get('storage')) #transform documents into tf-idf matrix
			click.secho('\ntfidf matrix generated\n',fg='green')
			vocab = vocab_build(config.get('storage')) #set of unique words from documents
			features = tf.get_feature_names() #the features (words) that tfidf vectorizer extracted
			features = set(features)
			vocab = vocab - (vocab - features) #get rid of discrepancies
			if vocab == features:
				click.secho("Vocab matches tfidf feature array\n", fg='green')
			vocab = sorted(list(vocab))
			tf_arr = tf_mat.toarray() 
			storage_weighted = attach_tfidf_weights(config.get('storage'), vocab, tf_arr) #attaches tf-idf weighting to each words
			#calculates vector representation of document as weighted average of word2vec vectors * tf-idf weights
			feature_vecs = featureVec(storage_weighted, model, 300)  
			sim_mat = compare(config.get('storage'), feature_vecs) # get cosine_similarity between each image pairing
			with open(pickle_filepath, 'w') as pickle_handle: #save the sim_mat in pickles/pickle_INPUT
				pickle.dump(sim_mat, pickle_handle)
		else: #just use word2vec
			feature_vecs_un = featureVec_unweighted(config.storage, model, 300)
			sim_mat = compare(config.storage, feature_vecs_un)
			with open(pickle_filepath, 'w') as pickle_handle:
				pickle.dump(sim_mat, pickle_handle)
	json_mat = json.dumps(sim_mat, cls=NumpyEncoder)
	config['sim_mat'] = json_mat #save sim_mat for outputs
	config.save()
	click.secho('saved word2vec sim_mat', fg='green')
	click.secho('Now run output to generate visualizations!', fg='red', blink=False)

@cli.command('output', short_help= click.style('Generate visual output', fg='red'))
@click.option('--no_thumb', is_flag=True, default = True, help = click.style("Don't draw the image thumbnails on iden_mat or con_mat(if images are unavailable)", fg='magenta'))
@click.option('--iden', is_flag=True, default = False, help = click.style("Generate an identity matrix with ski-kit learn manifold", fg='red'))
@click.option('--con', is_flag=True, default = False, help = click.style("Generate a confusion matrixMust include a --sep='index' flag", fg='magenta'))
@click.option('--sep', default = 1, help= click.style("Required for --con! sep=index separating two categories in the document/image ordering.", fg='red'))
@click.option('--spring', is_flag=True, default = False, help = click.style("Generate a graph in spring layout with pyplot", fg='magenta'))
@click.option('--mds', is_flag=True, default = False, help = click.style("Generate an MDS plot", fg='red'))
@click.option('--vis', is_flag=True, default= False, help = click.style('Write results to "nodes.txt" and "edges.txt" in vis.js graph formatThen you can manually copy into graph_vis.html', fg='magenta'))
@click.option('--explore', is_flag=True, default=False, help= click.style('Explore how the spring graph changes over a range of thresholds, specified by --thresh1=X and --thresh2=Y', fg='red'))
@click.option('--thresh1', default=0.6, help= click.style("Threshold to use for spring graphDefault = 0.6Also sets start threshold for --explore option", fg='magenta'))
@click.option('--thresh2', default=0.7, help= click.style("Threshold to use for vis.js outputDefault = 0.7Also sets the end threshold for --explore option", fg='red'))
@click.option('--step', default=0.01, help= click.style("Step size for explore optionDefault = 0.1", fg='magenta'))
@pass_config
def output(config, spring, mds, iden, con, sep, vis, no_thumb, explore, thresh1, thresh2, step):
	"""Generate visual outputs after running tfidf or word2vec"""
	json_mat = config.get('sim_mat')
	results = json.loads(json_mat, object_hook=json_numpy)
	if iden:
		click.secho('\ngenerating identity matrix output at iden_mat.png\n', fg='red', blink=False)
		clus = cluster(results, config.get('images'), False) #hierarchal centroid clustering
		clusmat = clus[0]
		clusim = clus[1] #new image ordering
		draw_id(clusmat, clusim, no_thumb) 
	if con:
		slices = conslice(results, sep) 
		clus = cluster(slices, config.get('images'), True)
		clusmat = clus[0]
		clusim = clus[1]
		click.secho('generating confusion matrix output at con_mat.png\n', fg='blue', blink=False)
		draw_con(clusmat, config.get('images'), clusim, sep)
	if spring: 
		click.secho('displaying spring graph\n', fg='yellow', blink=False)
		draw_spring_graph(results, thresh1, sep)
	if mds:
		click.secho('displaying MDS plot\n', fg='magenta', blink=False)
		draw_mds(results, sep)
	if vis:
		res_vis(results, config.get('images'), thresh2)
		click.secho('vis.js formatted graph sent to nodes.txt and edges.txt\n', fg='green')
	if explore:
		explore(results, thresh1, thresh2, step)
		click.secho('Saving exploration to explore_results', fg='green', blink=False)

def read_csv(path):
	"""Input a two column csv. Groups second column 
	entries based on matching texts in the first column
	 """
	answers = [] #descriptions
	images = [] #image titles
	with open(path, 'rU') as csvfile:
		reader = csv.reader(csvfile)
		index = 0
		for row in reader:
			im = (row[0])
			bothanswers = (row[1])
			split = bothanswers.split('|') #splits two mturk responses
			if im not in images:
				images.append(im)
				answers.append([[split[0]],[split[1]]])
				index += 1
			else:
				answers[index-1][0].append(split[0]) 
				answers[index-1][1].append(split[1])
	out = []
	out.append(images)
	out.append(answers)
	return out 

#reads .txt folders from specified folders
def read_folder(path):
	"""Reads folder of .txt files"""
	filenames = []
	answers = []
	for filename in glob.glob(os.path.join(path, '*.txt')):
		filenames.append(filename)
		f = open(filename, 'r')
		flines = f.readlines()
		descrps = ""
		for line in flines:		
			descrps += line
		answers.append(descrps)
	out = []
	out.append(filenames)
	out.append(answers)
	return out

#cleans mturk format text, gets specified POS, optional des or rem, optional stemming
def read_words(answers, rem, des, stem):
	"""Cleans text and does NLTK POS-tagging. Keeps NN, JJ, and RB.
	Optional stopword filtering and snowball stemming"
	Commented lines return total wordcount and counts for POS groups
	"""
	storage = []
	# wcount = [] #useful for getting statistics on words per response, parts of speech data
	# pos = [[],[],[]]
	# stopwords = nltk.corpus.stopwords.words('english') shouldn't effect noun, adj, and adv
	stemmer = SnowballStemmer("english") # 
	ind = [0,1]
	if rem == False: #default true. readin --rem option sets to False
		ind = [1]
	if des == False: #default true. readin --des options sets to False
		ind = [0]
	for i in range(len(answers)):
		sys.stdout.write(str(i)+",") #progress update for large input
		sys.stdout.flush()
		#ignore = ['black', 'white', 'image', 'picture'] #add your own set of ignore words to stopwords
		#stopwords += ignore
		descrps = ""
		for j in ind:
			for line in answers[i][j]:
				line = line.replace(',',' ')
				line = line.replace('.',' ')
				line = line.replace('-',' ')
				line = line.replace('/',' ')
				line = line.replace("'",' ')
				text = nltk.word_tokenize(line)
				text = nltk.pos_tag(text) #default maxent penn treebank
				wlist = ""
				for elem in text:
					# Each desired part of speech and not in ignored (not in use currently)
					if elem[1] in {'NN','NNS','NNP','NNPS','JJ','JJR','JJS','RB','RBR','RBS'}: #and elem[0] not in stopwords: 
						lower = (elem[0].lower())
						# if elem[1] in {'NN','NNS','NNP','NNPS'}:
						# 	pos[0].append(lower)
						# if elem[1] in {'JJ','JJR','JJS'}:
						# 	pos[1].append(lower)
						# if elem[1] in {'RB','RBR','RBS'}:
						# 	pos[2].append(lower)
						# wcount.append(lower)
						if stem: 
							lower = stemmer.stem(lower) #stemming for tfidf, won't for word2vec 
						wlist += lower
						wlist += " "
				descrps += wlist
		storage.append(descrps)
	return storage #wcount, pos]

#writes clean version of original text, selected POS, and counts for each POS
#configured for mturk format, lists descriptions them reminded
def clean_write(answers):
	"""Writes answers to word tables, original text, 
	POS-filtered words, Counts for each POS """
    document = Document() #microsoft word format
    document.add_heading('Listings Responses and Word Counts', 0)
    document.add_paragraph(
    'Image Descriptions and Reminded', 'ListBullet')
    #def tag_count(answers):
    image_counts = []
    for i in range(len(answers)):
        image_counts.append(Counter({}))
        document.add_paragraph(
        'Image Descriptions and Reminded', style='ListBullet')
        for j in [0,1]: #do 
            table = document.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Cleaned Response Text'
            hdr_cells[1].text = 'Usable Words'
            hdr_cells[2].text = 'Word Counts'
            for k in range(len(answers[i][j])):
                row_cells = table.add_row().cells
                txt = answers[i][j][k]
                txt = txt.replace(',','')
                txt = txt.replace('.','')
                txt = txt.replace('-',' ')
                txt = txt.replace('/',' ')
                txt=txt.lower()
                row_cells[0].text = txt
                words = nltk.word_tokenize(txt)
                tags= nltk.pos_tag(words)
                word_bag = []
                tag_bag = []
                keep_tags = {'NN','NNS','NNP','NNPS','JJ','JJR','JJS','RB','RBR','RBS'}
                for tag in tags:
                    if tag[1] in keep_tags:
                        word_bag.append(tag[0])
                        tag_bag.append(tag[1])
                row_cells[1].text = str(word_bag)
                counts = Counter(tag_bag)
                POS = Counter(tag_bag).most_common(10)
                image_counts[i] += counts 
                row_cells[2].text = str(POS)
                answers[i][j][k] = word_bag  
    document.save('response_&_wordcount.docx')

def vocab_build(corpus):
	"""Builds set of unique words """
	lexicon = set()
	for doc in corpus:
		doc = doc.split()
		lexicon.update([word for word in doc])
	return lexicon

def attach_tfidf_weights(storage, vocab, tf_arr):
	"""Appends tf-idf weights to each word """
	wordlist = vocab
	storage_weighted = []
	for i in range(len(storage)):
		sys.stdout.write(str(i)+",")
		sys.stdout.flush()
		docweights = []
		stor_list = storage[i].split()
		for word in stor_list:
			words = [word,0]
			for j in range(len(wordlist)):
				if (wordlist[j] == word):
					words[1] = tf_arr[i][j]	
			docweights.append(words)
		storage_weighted.append(docweights)
	return storage_weighted

def featureVec(storage, model, num_features):
	"""creates a vector representation for each image's descriptions (document)
	"""
	index2word_set = set(model.index2word)
	realWords = [] #in model
	notfound = [] 
	feature_vecs = []
	tot_wei = 0.0 #tf-idf weight total
	for i in range(len(storage)):
		realWords.append([])
		for word in storage[i]:
			#cap = word[0].capitalize() catch if capitalized proper noun in model
			if (word[0] in index2word_set):
				realWords[i].append(word)
				tot_wei += word[1]
				continue
	print tot_wei
	for i in range(len(realWords)):
		feature_vec = np.zeros((num_features), dtype="float32")
		num_words = 0
		for realword in realWords[i]:
			weighted_vec = model[realword[0]]*(realword[1] / tot_wei) #normalized tf-idf weight
			feature_vec = np.add(feature_vec, weighted_vec)
			num_words += 1
		feature_vec = np.divide(feature_vec, num_words) #average of each word vector
		feature_vecs.append(feature_vec)
	return feature_vecs

def featureVec_unweighted(storage, model, num_features):
	""" Same as featureVec, but no tf-idf weights"""
	index2word_set = set(model.index2word)
	realWords = []
	feature_vecs = []
	for i in range(len(storage)):
		realWords.append([])
		storage[i] = storage[i].split()
		for word in storage[i]:
			if word in index2word_set:
				realWords[i].append(word)
			else:
				click.secho("notfound:  ", fg='red')
				click.echo(word)
	for i in range(len(realWords)):
		feature_vec = np.zeros((num_features), dtype="float32")
		num_words = 0
		for realword in realWords[i]:
			weighted_vec = model[realword]
			feature_vec = np.add(feature_vec, weighted_vec)
			num_words += 1
		feature_vec = np.divide(feature_vec, num_words)
		feature_vecs.append(feature_vec)
	return feature_vecs

def compare(storage, feature_vecs):
	results = zeros((len(storage),len(storage)))
	min_result = 1.0
	max_result = 0.0 
	for i in range(len(storage)):
		for j in range(len(storage)):
			result = cosine_similarity(feature_vecs[i], feature_vecs[j])
			results[i][j] = result
			if result < min_result:
				min_result = result
			if result > max_result:
				max_result = result
	# 	sys.stdout.write('.') #progress
	# 	sys.stdout.flush()
	# sys.stdout.write('\n')
	#used normalize similarity scores from 0 to 1
	print 'max: ' + str(max_result)
	print 'min: ' + str(min_result)
	max_result -= min_result
	for i in range(len(results)): #normalization 
		for j in range(len(results[i])):
			results[i][j] = (results[i][j] - min_result) / max_result 
	return results

def draw_id(results, stored_images, thumb):
	"""draws an indentity matrix from the similarity matrix"""
	stored = []
	for filename in stored_images:
		stored.append("mturk_images/"+filename) #edit to match unique image folder name
	stored_images1 = stored
	stored_images2 = stored[::-1] #reverse for x-axis thumbs
	im = Image.new('RGB', ((len(stored_images2)+2)*50, (len(stored_images1)+2)*50), (255, 255, 255))
	draw = ImageDraw.Draw(im)
	if thumb == True:
		click.secho('drawing thumbnails', fg='red', blink=False)
		for i in range(len(stored_images2)):
			sys.stdout.write(".")
			sys.stdout.flush()
			thumb = Image.open(stored_images2[i]) #stored_images_for_ident
			thumb = thumb.resize((50, 50), Image.ANTIALIAS)
			im.paste(thumb, box=(50*i+50, 0, 50*i+100, 50)) #box position =(x1, y1, x2, y2) from top left
			im.paste(thumb, box=(50*i+50, 50*len(stored_images2)+50, 50*i+100, 50*len(stored_images2)+100)) 
			#draw on both sides for ease of viewing
		sys.stdout.write('\n')
		sys.stdout.flush()
		for j in range(len(stored_images1)):
			sys.stdout.write(".")
			sys.stdout.flush()
			thumb = Image.open(stored_images1[j])
			thumb = thumb.resize((50, 50), Image.ANTIALIAS)
			im.paste(thumb, box=(0, 50*j+50, 50, 50*j+100)) 
			im.paste(thumb, box=(50*len(stored_images1)+50, 50*j+50, 50*len(stored_images1)+100, 50*j+100))
		sys.stdout.write('\n')
		sys.stdout.flush()	
	# Draw color squares
	for j in range(len(results)):
		for k in range(len(results)):
			avg_score = results[j][k]
			r_amount = 0
			g_amount = 0
			b_amount = 0
			if avg_score >= 0.75: # yellow to red
				r_amount = 255
				g_amount = 255 - 255*4*(avg_score-0.75)
			if avg_score >= 0.5 and avg_score <= 0.75: #green to yellow
				r_amount = 255*4*(avg_score-0.5)
				g_amount = 255
			if avg_score >= 0.25 and avg_score <= 0.5: #cyan to green
				g_amount = 255
				b_amount = 255 - 255*4*(avg_score-0.25)
			if avg_score <= 0.25: #blue to cyan
				g_amount = 255*4*(avg_score)
				b_amount = 255
			score = round(avg_score, 3)
			score = str(score)
			draw.rectangle([k*50+50, j*50+50, k*50+100, j*50+100], fill=(int(r_amount), int(g_amount), int(b_amount)), outline=None)
			draw.text((k*50+58, j*50+80), score, fill=(0,0,0))
	del draw
	im.save('iden_mat.png', "PNG")

def draw_con(results, stored_images, clusim, sep):
	"""draw confusion matrix separating two categories, determined by sep"""
	#normalize confusion matrix
	images = clusim 
	scenes = stored_images[sep:]
	conmax = np.amax(results) #normalization setup
	conmin = np.amin(results)
	conmax = conmax - conmin
	thumbs_x = [""]*len(scenes)
	for i in range(len(scenes)):
		thumbs_x[i] = "mturk_images/"+scenes[i]
	thumbs_y = [""]*len(images)
	for j in range(len(images)):
		thumbs_y[j] = "mturk_images/"+images[j]
	dim = shape(results) 
	x = dim[1]
	y = dim[0]
	im = Image.new('RGB', ((x+2)*50, (y+2)*50), (255, 255, 255))
	draw = ImageDraw.Draw(im)
	click.secho('drawing thumbnails', fg='blue', blink=False)
	for j in range(y):
		sys.stdout.write(".")
		sys.stdout.flush()
		thumb = Image.open(thumbs_y[j])
		thumb = thumb.resize((50, 50), Image.ANTIALIAS)
		im.paste(thumb, box=(0, 50*j+50, 50, 50*j+100))
		im.paste(thumb, box=(50*x+50, 50*j+50, 50*x+100, 50*j+100))
	sys.stdout.write('\n')
	sys.stdout.flush()
	for i in range(x):
		sys.stdout.write(".")
		sys.stdout.flush()
		thumb = Image.open(thumbs_x[i]) #images_for_ident
		thumb = thumb.resize((50, 50), Image.ANTIALIAS)
		im.paste(thumb, box=(50*i+50, 0, 50*i+100, 50))
		im.paste(thumb, box=(50*i+50, 50*y+50, 50*i+100, 50*y+100))
	sys.stdout.write('\n')
	sys.stdout.flush()
	#Draw color squares
	#Images and scenes tend to have lower similarities
	#Colors are adjusted to (0, .2, .4, .6) scale to compensate
	for j in range(y):
		for k in range(x):
			avg_score = results[j][k]
			avg_score = (avg_score - conmin) / conmax
			r_amount = 0
			g_amount = 0
			b_amount = 0
			if avg_score >= 0.6:
				r_amount = 255
				g_amount = 255 - 255*4*(avg_score-0.6)
			if avg_score >= 0.4 and avg_score <= 0.6:
				r_amount = 255*5*(avg_score-0.4)
				g_amount = 255
			if avg_score >= 0.2 and avg_score <= 0.4:
				g_amount = 255
				b_amount = 255 - 255*5*(avg_score-0.2)
			if avg_score <= 0.2:
				g_amount = 255*4*(avg_score)
				b_amount = 255
			score = round(avg_score, 3)
			score = str(score)
			draw.rectangle([k*50+50, j*50+50, k*50+100, j*50+100], fill=(int(r_amount), int(g_amount), int(b_amount)), outline=None)
			draw.text((k*50+58, j*50+80), score, fill=(0,0,0))
	del draw
	im.save('con_mat.png')

def draw_spring_graph(results, thresh, sep):
	"""Draw graph in spring layoout, 
	force-directed algorithm puts similar image nodes close to eachother
	Assumes symmetric split with the two categories (artificial/natural + indoor/outdoor for ComCon)"""
	img_colors = ['green'] * (sep/2) #add or modify indices to get new colors 
	img_colors2 = ['red'] * (sep/2)
	sce_colors = ['blue'] * ((len(results) - sep)/2)
	sce_colors2 = ['yellow'] * ((len(results) - sep)/2)
	node_colors = img_colors + img_colors2 + sce_colors +sce_colors2
	res2 = np.copy(results)
	low_val = res2 < thresh
	res2[low_val] = 0
	graph = nx.from_numpy_matrix(res2)
	pos = nx.spring_layout(graph)
	nx.draw_networkx_nodes(graph, pos=pos, node_color = node_colors)
	nx.draw_networkx_edges(graph, pos=pos)
	# xs = [] # Add labels (looks pretty messy with large graph)
	# ys = []
	# for i in range(len(pos)):
	# 	xs.append(pos[i][0])
	# 	ys.append(pos[i][1])
	# for label, x, y, in zip(names, xs, ys):
	# 	plt.annotate(
	# 	label,
	# 	xy = (x, y), xytext = (-10, 35),
	# 	textcoords = 'offset points', ha = 'right', va = 'bottom',
	# 	bbox = dict(boxstyle = 'round,pad=0.5', fc = 'green', alpha = 0.7),
		# arrowprops = dict(arrowstyle = '->', connectionstyle = 'arc3,rad=0'))
	plt.show()


def draw_mds(results, sep):
	"""draws MDS plot of high-dimension vectors"""
	mds = manifold.MDS(n_components=2, dissimilarity="precomputed", random_state=6)
	a = (sep/2) #modify indices to match data
	b = sep
	c = ((len(results)-sep)/2)+b
	res2 = np.copy(results)
	res2 = 1 - results
	result = mds.fit(res2)
	coords = result.embedding_
	plt.subplots_adjust(bottom = 0.1)
	plt.scatter(
		coords[:a, 0], coords[:a, 1], color = 'green', marker = 'o', s=70)
	plt.scatter(
		coords[a:b, 0], coords[a:b, 1], color = 'red', marker = 'o', s=70)
	plt.scatter(
		coords[b:c, 0], coords[b:c, 1], color = 'blue', marker = 'v', s=70)
	plt.scatter(
		coords[c:, 0], coords[c:, 1], color = 'orange', marker = 'v', s=70)
	# for label, x, y in zip(names, coords[:, 0], coords[:, 1]):
	#     plt.annotate(
	#         label,
	#         xy = (x, y), xytext = (-10, 25),
	#         textcoords = 'offset points', ha = 'right', va = 'bottom',
	#         bbox = dict(boxstyle = 'round,pad=0.5', fc = 'green', alpha = 0.7),
	#         arrowprops = dict(arrowstyle = '->', connectionstyle = 'arc3,rad=0'))
	plt.show()


def res_vis(results, images, thresh):
	"""Saves results in vis.js formatted .txt files. 
	After generating, manually copy nodes.txt and edges.txt into graph_vis.html template,
	then delete .txt files
	"""
	res2 = np.copy(results)
	nodes = []
	edges = []
	for i in range(len(results)):
		#assumes html will be stored one directory below image folder
		string = "{id: " +str(i)+ ", image: '../mturk_images/" +images[i]+ "'},"
		nodes.append(string)
	for i in range(len(results)):
		for j in range(i+1, len(results)):
			if res2[i][j] >= thresh:
				string = "{from: "+str(i)+", to: "+str(j)+"},"
				edges.append(string)
	f=open('nodes.txt', 'w')
	for node in nodes:
		f.write(node+'\n')
	f=open('edges.txt', 'w')
	for edge in edges:
		f.write(edge+'\n')

def explorate(results, thresh1, thresh2, step, sep):
	"""view spring graphs and degree distributions across a range of thresholds
	Default is plt.show() for each step, uncomment last lines to save figures in explore_results folder
	"""
	i = 0
	img_colors = ['green'] * (sep/2) #indices separating different node colors
	img_colors2 = ['red'] * (sep/2)
	sce_colors = ['blue'] * ((len(results) - sep)/2)
	sce_colors2 = ['yellow'] * ((len(results) - sep)/2)
	node_colors = img_colors + img_colors2 + sce_colors +sce_colors2
	for thresh in np.arange(thresh1, thresh2, step):
		print thresh
		res2 = np.copy(results)
		hi_val = (res2 >= thresh)
		graph = nx.from_numpy_matrix(hi_val)
		pos = nx.spring_layout(graph)
		nx.draw_networkx_nodes(graph, pos=pos, node_color = node_colors)
		nx.draw_networkx_edges(graph, pos=pos)
 		plt.show()
 		# plt.savefig('explore_results/spring'+str(i)+'.png')
 		#plt.close('all')
 		i+=1

 	i = 0
	for thresh in np.arange(thresh1, thresh2, step):
		print thresh
		res2 = np.copy(results)
		hi_val = (res2 >= thresh)
		graph = nx.from_numpy_matrix(hi_val)		
		degrees = graph.degree()
		values = sorted(set(degrees.values()))
		hist = [degrees.values().count(x) for x in values]

		plt.figure()
		plt.grid(True)
		plt.plot(values, hist, 'ro-') 
		plt.legend(['degree'])
		plt.xlabel('Degree')
		plt.ylabel('Number of nodes')
		plt.title('MTurk Degree Distribution')
		plt.xlim([0, 35])
		plt.ylim([0, 25])
		plt.show()
		# plt.savefig('explore_results/distribution'+str(i)+'.png')
		# plt.close('all')
		i += 1

def conslice(sim_mat, sep):
	"""Slices a confusion matrix out of similarity matrix based on sep"""
	images = sim_mat[:sep]
	slices = []
	for i in range(len(images)):
		slices.append(images[i][sep:])
	return slices 

def cluster(sim_mat, images, con):
	"""hierachal clustering"""
	if con: #only cluster along y-axis
		dim = shape(sim_mat)
		x = dim[0]
		images = images[:x]
	df = pd.DataFrame(sim_mat)
	hier = hc.linkage(sim_mat, method='centroid')
	lev = hc.leaves_list(hier)
	# get the clustered indices
	mat = df.iloc[lev,:]
	if not con:
		mat = mat.iloc[:, lev[::-1]]
	#get hierarchal indices back to numpy matrix
	simul = mat.as_matrix()
	#to get new image list assuming original list is ims
	imin = mat.index.values
	clusim = [""]*len(images)
	for i in range(len(images)):
		clusim[i] = images[imin[i]]
	return [simul, clusim]







