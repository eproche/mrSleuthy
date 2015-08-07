import nltk, glob, os, sys
import csv
from gensim.models import word2vec
from nltk.stem.snowball import SnowballStemmer
from docx import Document
from collections import Counter 

def read_csv(path):
	"""Input a two column csv. Groups second column 
	entries based on matching texts in the first column
	 """
	answers = [] #descriptions
	images = [] #image titles
	with open(path, 'rU') as csvfile:
		reader = csv.reader(csvfile)
		index = 0
		for row in reader:
			im = (row[0])
			bothanswers = (row[1])
			split = bothanswers.split('|') #splits two mturk responses
			if im not in images:
				images.append(im)
				answers.append([[split[0]],[split[1]]])
				index += 1
			else:
				answers[index-1][0].append(split[0]) 
				answers[index-1][1].append(split[1])
	out = []
	out.append(images)
	out.append(answers)
	return out 

#reads .txt folders from specified folders
def read_folder(path):
	"""Reads folder of .txt files"""
	filenames = []
	answers = []
	for filename in glob.glob(os.path.join(path, '*.txt')):
		filenames.append(filename)
		f = open(filename, 'r')
		flines = f.readlines()
		descrps = ""
		for line in flines:		
			descrps += line
		answers.append(descrps)
	out = []
	out.append(filenames)
	out.append(answers)
	return out

#cleans mturk format text, gets specified POS, optional des or rem, optional stemming
def read_words(answers, rem, des, stem):
	"""Cleans text and does NLTK POS-tagging. Keeps NN, JJ, and RB.
	Optional stopword filtering and snowball stemming"
	Commented lines return total wordcount and counts for POS groups
	"""
	storage = []
	# wcount = [] #useful for getting statistics on words per response, parts of speech data
	# pos = [[],[],[]]
	# stopwords = nltk.corpus.stopwords.words('english') shouldn't effect noun, adj, and adv
	stemmer = SnowballStemmer("english") # 
	ind = [0,1]
	if rem == False: #default true. readin --rem option sets to False
		ind = [1]
	if des == False: #default true. readin --des options sets to False
		ind = [0]
	for i in range(len(answers)):
		sys.stdout.write(str(i)+",") #progress update for large input
		sys.stdout.flush()
		#ignore = ['black', 'white', 'image', 'picture'] #add your own set of ignore words to stopwords
		#stopwords += ignore
		descrps = ""
		for j in ind:
			for line in answers[i][j]:
				line = line.replace(',',' ')
				line = line.replace('.',' ')
				line = line.replace('-',' ')
				line = line.replace('/',' ')
				line = line.replace("'",' ')
				text = nltk.word_tokenize(line)
				text = nltk.pos_tag(text) #default maxent penn treebank
				wlist = ""
				for elem in text:
					# Each desired part of speech and not in ignored (not in use currently)
					if elem[1] in {'NN','NNS','NNP','NNPS','JJ','JJR','JJS','RB','RBR','RBS'}: #and elem[0] not in stopwords: 
						lower = (elem[0].lower())
						# if elem[1] in {'NN','NNS','NNP','NNPS'}:
						# 	pos[0].append(lower)
						# if elem[1] in {'JJ','JJR','JJS'}:
						# 	pos[1].append(lower)
						# if elem[1] in {'RB','RBR','RBS'}:
						# 	pos[2].append(lower)
						# wcount.append(lower)
						if stem: 
							lower = stemmer.stem(lower) #stemming for tfidf, won't for word2vec 
						wlist += lower
						wlist += " "
				descrps += wlist
		storage.append(descrps)
	return storage #wcount, pos]

# def read_words(answers):
# 	"""Cleans text and does NLTK POS-tagging. Keeps NN, JJ, and RB.
# 	Optional stopword filtering and snowball stemming"
# 	Commented lines return total wordcount and counts for POS groups
# 	"""
# 	storage = []
# 	for line in answers:
# 		line = line.replace(',',' ')
# 		line = line.replace('.',' ')
# 		line = line.replace('-',' ')
# 		line = line.replace('/',' ')
# 		line = line.replace("'",' ')
# 		print line
# 		text = nltk.word_tokenize(line)
# 		text = nltk.pos_tag(text) #default maxent penn treebank
# 		wlist = ""
# 		for elem in text:
# 			# Each desired part of speech and not in ignored (not in use currently)
# 			if elem[1] in {'NN','NNS','NNP','NNPS','JJ','JJR','JJS','RB','RBR','RBS'}: #and elem[0] not in stopwords: 
# 				lower = (elem[0].lower())
# 				wlist += lower
# 				wlist += " "
# 		storage.append(wlist)
# 	return storage #wcount, pos]

#writes clean version of original text, selected POS, and counts for each POS
#configured for mturk format, lists descriptions them reminded
def clean_write(answers):
    """Writes answers to word tables, original text, 
    POS-filtered words, Counts for each POS """
    document = Document() #microsoft word format
    document.add_heading('Listings Responses and Word Counts', 0)
    document.add_paragraph(
    'Image Descriptions and Reminded', 'ListBullet')
    #def tag_count(answers):
    image_counts = []
    for i in range(len(answers)):
        image_counts.append(Counter({}))
        document.add_paragraph(
        'Image Descriptions and Reminded', style='ListBullet')
        for j in [0,1]: #do 
            table = document.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Cleaned Response Text'
            hdr_cells[1].text = 'Usable Words'
            hdr_cells[2].text = 'Word Counts'
            for k in range(len(answers[i][j])):
                row_cells = table.add_row().cells
                txt = answers[i][j][k]
                txt = txt.replace(',','')
                txt = txt.replace('.','')
                txt = txt.replace('-',' ')
                txt = txt.replace('/',' ')
                txt=txt.lower()
                row_cells[0].text = txt
                words = nltk.word_tokenize(txt)
                tags= nltk.pos_tag(words)
                word_bag = []
                tag_bag = []
                keep_tags = {'NN','NNS','NNP','NNPS','JJ','JJR','JJS','RB','RBR','RBS'}
                for tag in tags:
                    if tag[1] in keep_tags:
                        word_bag.append(tag[0])
                        tag_bag.append(tag[1])
                row_cells[1].text = str(word_bag)
                counts = Counter(tag_bag)
                POS = Counter(tag_bag).most_common(10)
                image_counts[i] += counts 
                row_cells[2].text = str(POS)
                answers[i][j][k] = word_bag  
    document.save('response_&_wordcount.docx')

    